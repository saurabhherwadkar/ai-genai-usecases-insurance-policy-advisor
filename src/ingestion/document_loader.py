"""
Document loader module for the Insurance Policy Advisor.

Handles loading of insurance policy documents in multiple formats
(PDF, DOCX, TXT). Each format has a dedicated loading method that
extracts raw text content from the file.
"""

from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field

from src.utils.logger import get_logger

# Module-level logger for document loading operations
logger = get_logger(__name__)


class Document(BaseModel):
    """
    Represents a loaded document with its content and metadata.

    Attributes:
        content: The raw text content extracted from the document.
        metadata: Dictionary containing document metadata like filename, format, and path.
    """

    content: str = Field(description="The raw text content of the document")
    metadata: dict = Field(default_factory=dict, description="Document metadata including source file info")


class DocumentLoader:
    """
    Loads insurance policy documents from various file formats.

    Supports PDF, DOCX, and TXT formats. Each format has a dedicated
    method that handles the specifics of text extraction for that type.
    """

    def __init__(self) -> None:
        """Initialize the DocumentLoader with format-to-method mapping."""
        # Map file extensions to their corresponding loader methods
        self._loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_txt,
        }

    def load(self, file_path: str) -> Optional[Document]:
        """
        Load a document from the given file path.

        Determines the file format from its extension and delegates
        to the appropriate loader method.

        Args:
            file_path: Path to the document file to load.

        Returns:
            A Document object containing the extracted text and metadata,
            or None if the file format is unsupported or loading fails.
        """
        path = Path(file_path)

        # Validate that the file exists
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        # Get the file extension in lowercase for format matching
        extension = path.suffix.lower()

        # Check if the format is supported
        if extension not in self._loaders:
            logger.warning(f"Unsupported file format: {extension} for file: {file_path}")
            return None

        try:
            # Delegate to the format-specific loader method
            logger.info(f"Loading document: {path.name} (format: {extension})")
            content = self._loaders[extension](path)

            # Create and return the Document with metadata
            document = Document(
                content=content,
                metadata={
                    "filename": path.name,
                    "format": extension,
                    "path": str(path.absolute()),
                    "size_bytes": path.stat().st_size,
                },
            )

            logger.info(f"Successfully loaded document: {path.name} ({len(content)} characters)")
            return document

        except Exception as error:
            logger.error(f"Failed to load document {file_path}: {str(error)}")
            return None

    def load_directory(self, directory_path: str, supported_formats: list[str] = None) -> list[Document]:
        """
        Load all supported documents from a directory.

        Scans the directory for files with supported extensions and
        loads each one using the appropriate loader method.

        Args:
            directory_path: Path to the directory containing documents.
            supported_formats: List of file extensions to load. Defaults to all supported formats.

        Returns:
            A list of successfully loaded Document objects.
        """
        directory = Path(directory_path)

        # Validate that the directory exists
        if not directory.exists() or not directory.is_dir():
            logger.error(f"Directory not found or not a directory: {directory_path}")
            return []

        # Use all supported formats if none specified
        formats = supported_formats or list(self._loaders.keys())

        # Collect all files matching supported formats
        documents = []
        for format_ext in formats:
            # Find all files with this extension in the directory
            for file_path in directory.glob(f"*{format_ext}"):
                document = self.load(str(file_path))
                if document is not None:
                    documents.append(document)

        logger.info(f"Loaded {len(documents)} documents from directory: {directory_path}")
        return documents

    def _load_pdf(self, path: Path) -> str:
        """
        Extract text content from a PDF file.

        Uses PyPDF2 to read each page and concatenate the text.

        Args:
            path: Path object pointing to the PDF file.

        Returns:
            The extracted text content from all pages.
        """
        from PyPDF2 import PdfReader

        # Open and read the PDF file
        reader = PdfReader(str(path))

        # Extract text from each page and join with newlines
        pages_text = []
        for page_number, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
                logger.debug(f"Extracted text from page {page_number + 1} of {path.name}")

        return "\n\n".join(pages_text)

    def _load_docx(self, path: Path) -> str:
        """
        Extract text content from a DOCX file.

        Uses python-docx to read paragraphs from the Word document.

        Args:
            path: Path object pointing to the DOCX file.

        Returns:
            The extracted text content from all paragraphs.
        """
        from docx import Document as DocxDocument

        # Open and read the DOCX file
        doc = DocxDocument(str(path))

        # Extract text from each paragraph
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        logger.debug(f"Extracted {len(paragraphs)} paragraphs from {path.name}")
        return "\n\n".join(paragraphs)

    def _load_txt(self, path: Path) -> str:
        """
        Read text content from a plain text file.

        Args:
            path: Path object pointing to the TXT file.

        Returns:
            The raw text content of the file.
        """
        # Read the file content with UTF-8 encoding
        with open(path, "r", encoding="utf-8") as text_file:
            content = text_file.read()

        logger.debug(f"Read {len(content)} characters from {path.name}")
        return content
